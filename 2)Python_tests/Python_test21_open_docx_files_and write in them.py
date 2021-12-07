from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt

#init
document = Document()
run = document.add_paragraph().add_run()
font = run.font

#heading
document.add_heading('Document Title', 0)

#first paragraph with bold and italic
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ') 
p.add_run('italic.').italic = True

#second paragraph with red and size 22
wp = document.add_paragraph()
wp = wp.add_run('I want this sentence colored red with fontsize=22')
wp.font.size = Pt(22)
wp.font.color.rgb = RGBColor(255,0,0)

#second heading
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

#third paragraph
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
#fourth paragraph
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

#picture
document.add_picture('discord.JPG', width=Inches(1.25))

#table
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

#end of the page
document.add_page_break()

#save
document.save('C:/Users/danie/Desktop/Tests/demo.docx')
