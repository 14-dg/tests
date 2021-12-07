
from docx import Document
document = Document('C:/Users/danie/Desktop/Tests/demo.docx')

# document.paragraphs  # to extract paragraphs
# document.paragraphs[2].text  # gives the text
# ​
for par in document.paragraphs:  # to extract the whole text
  print(par.text)


i=0
for par in document.paragraphs:  # to extract the whole text
    i+=1
    if 'I' in par.text:
        print(par.text)
        print(i)